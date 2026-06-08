import os
import io
import zipfile
import json
import uuid
import datetime
import hashlib
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, session
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from pypdf import PdfReader
import openai

app = Flask(__name__, template_folder='templates', static_folder='static')
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "STORMCORE_SOVEREIGN_2026_MASTERKEY_X")

# System Authentication Matrices
openai.api_key = os.environ.get("OPENAI_API_KEY", "mock-key")

# Core Database Mock Layers (Persistent state mapping for 2026 multi-tier engines)
SYSTEM_DB = {
    "users": {
        "cust_1": {"role": "customer", "name": "John Doe", "tier": "free", "cv_text": "", "job_description": "", "optimized_cv": "", "cover_letter": "", "qa_prep": [], "usage_count": 0},
        "admin_1": {"role": "admin", "name": "System Administrator"},
        "owner_1": {"role": "owner", "name": "Stormcore Director"}
    },
    "transactions": [],
    "marketing_metrics": {
        "daily_signups_basic": 42, 
        "total_ad_spend_zar": 12450.00,
        "conversions_today": 48,
        "impulse_clicks": 1420
    },
    "financial_ledgers": {
        "owner_standard_bank_total_zar": 125000.00,
        "admin_african_bank_total_zar": 25000.00,
        "company_registration_pool_zar": 4850.00,  # Maxes out at R5000, then shifts to hosting
        "hosting_fees_pool_zar": 0.00,
        "system_upgrades_pool_zar": 75000.00
    },
    "store_products": [
        {"id": "prod_1", "name": "Premium High-Trust Resume Rewrite Template Pack", "price": 150.00, "type": "once_off", "tier": "basic"},
        {"id": "prod_2", "name": "Pro Master Portfolio & LinkedIn Matrix Layouts", "price": 350.00, "type": "once_off", "tier": "pro"},
        {"id": "prod_3", "name": "Enterprise Career Launch Blueprint Suite", "price": 600.00, "type": "once_off", "tier": "enterprise"},
        {"id": "sub_1", "name": "Basic Career Tracker Subscription", "price": 120.00, "type": "subscription", "tier": "basic"},
        {"id": "sub_2", "name": "Pro Network Engagement Engine Subscription", "price": 300.00, "type": "subscription", "tier": "pro"},
        {"id": "sub_3", "name": "Enterprise Autopilot Executive Placement Subscription", "price": 550.00, "type": "subscription", "tier": "enterprise"},
        {"id": "sub_4", "name": "Sovereign Real-Time 2026 AI Coach Access Pack", "price": 450.00, "type": "subscription", "tier": "tutor"}
    ]
}

def get_session_user():
    if "user_id" not in session:
        session["user_id"] = "cust_1"
    uid = session["user_id"]
    return SYSTEM_DB["users"].get(uid, SYSTEM_DB["users"]["cust_1"])

# --- Document Processing Systems (Full File Reads & Outputs) ---
def extract_text_from_pdf(file_bytes):
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        return f"PDF Extractor Interruption: {str(e)}"

def generate_pdf_bytes(title, content):
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=20, leading=24, spaceAfter=15)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=8)
    
    story = [Paragraph(title, title_style), Spacer(1, 10)]
    for line in content.split('\n'):
        if line.strip():
            story.append(Paragraph(line.strip(), body_style))
    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()

def generate_docx_bytes(title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()

# --- Automated Ledger Payout Core Logic ---
def execute_ledger_split(amount_zar):
    ledger = SYSTEM_DB["financial_ledgers"]
    
    # 50% Owner Standard Bank, 10% Admin African Bank, 30% Upgrades
    split_owner = amount_zar * 0.50
    split_admin = amount_zar * 0.10
    split_upgrades = amount_zar * 0.30
    
    # 10% Company Registration Pool capping at R5000, then routing to hosting fees pool
    split_reg_or_hosting = amount_zar * 0.10
    
    ledger["owner_standard_bank_total_zar"] += split_owner
    ledger["admin_african_bank_total_zar"] += split_admin
    ledger["system_upgrades_pool_zar"] += split_upgrades
    
    if ledger["company_registration_pool_zar"] < 5000.00:
        ledger["company_registration_pool_zar"] += split_reg_or_hosting
        if ledger["company_registration_pool_zar"] > 5000.00:
            overflow = ledger["company_registration_pool_zar"] - 5000.00
            ledger["company_registration_pool_zar"] = 5000.00
            ledger["hosting_fees_pool_zar"] += overflow
    else:
        ledger["hosting_fees_pool_zar"] += split_reg_or_hosting

# --- Framework UI Navigation Render Routes ---
@app.route('/')
def index():
    return render_template('index.html', user=get_session_user())

@app.route('/store')
def store():
    return render_template('store.html', products=SYSTEM_DB["store_products"], user=get_session_user())

@app.route('/cart')
def cart():
    if "cart" not in session:
        session["cart"] = []
    cart_items = [p for p in SYSTEM_DB["store_products"] if p["id"] in session["cart"]]
    total = sum(item["price"] for item in cart_items)
    return render_template('cart.html', cart=cart_items, total=total, user=get_session_user())

@app.route('/dashboard/customer')
def dashboard_customer():
    return render_template('dashboard_customer.html', user=get_session_user())

@app.route('/dashboard/admin')
def dashboard_admin():
    return render_template('dashboard_admin.html', user=get_session_user(), metrics=SYSTEM_DB["marketing_metrics"])

@app.route('/dashboard/owner')
def dashboard_owner():
    return render_template('dashboard_owner.html', user=get_session_user(), ledgers=SYSTEM_DB["financial_ledgers"])

@app.route('/tutor')
def tutor():
    return render_template('tutor.html', user=get_session_user())

@app.route('/payment')
def payment():
    return render_template('payment.html', user=get_session_user())

@app.route('/marketing')
def marketing():
    return render_template('marketing.html', user=get_session_user())

# --- User Simulation Switching Node ---
@app.route('/api/switch_role/<role>')
def switch_role(role):
    if role == "admin":
        session["user_id"] = "admin_1"
    elif role == "owner":
        session["user_id"] = "owner_1"
    else:
        session["user_id"] = "cust_1"
    return redirect('/')

# --- Cart Operations Node ---
@app.route('/api/cart/add', methods=['POST'])
def add_to_cart():
    pid = request.json.get("product_id")
    if "cart" not in session:
        session["cart"] = []
    if pid not in session["cart"]:
        session["cart"].append(pid)
    return jsonify({"status": "success", "cart_count": len(session["cart"])})

@app.route('/api/cart/clear')
def clear_cart():
    session["cart"] = []
    return redirect('/cart')

# --- E-Commerce Processing & Live Payment Integration Matrix ---
@app.route('/api/payment/checkout', methods=['POST'])
def api_checkout_process():
    gateway = request.json.get("gateway", "direct_eft")
    amount = float(request.json.get("amount", 0.00))
    tier_upgrade = request.json.get("tier", "free")
    
    if amount <= 0:
        return jsonify({"status": "error", "message": "Transaction total value must exceed zero."}), 400
        
    tx_id = f"STORMCORE-TX-{uuid.uuid4().hex[:10].upper()}"
    
    # 2026 Core Live Integrated Signatures Builder Parameters
    if gateway == "payfast":
        merchant_id = os.environ.get("PAYFAST_MERCHANT_ID", "10000100")
        merchant_key = os.environ.get("PAYFAST_MERCHANT_KEY", "46f0cd694add8")
        pf_data = {
            "merchant_id": merchant_id, "merchant_key": merchant_key,
            "amount": f"{amount:.2f}", "item_name": "Stormcore Platform Purchase",
            "return_url": request.host_url + "dashboard/customer?status=success",
            "cancel_url": request.host_url + "cart?status=cancelled"
        }
        # Simulate redirection signature generation string
        query_string = "&".join([f"{k}={v}" for k, v in pf_data.items()])
        pf_url = f"https://sandbox.payfast.co.za/eng/process?{query_string}"
        execute_ledger_split(amount)
        if get_session_user()["role"] == "customer":
            SYSTEM_DB["users"]["cust_1"]["tier"] = tier_upgrade
        session["cart"] = []
        return jsonify({"status": "redirect", "url": pf_url})

    elif gateway == "ozow":
        site_code = os.environ.get("OZOW_SITE_CODE", "STORMCORE01")
        private_key = os.environ.get("OZOW_PRIVATE_KEY", "secure_key")
        bank_reference = tx_id
        hash_string = f"{site_code}{bank_reference}{amount:.2f}ZARhttps://stormcore.co.za{private_key}"
        generated_hash = hashlib.sha256(hash_string.encode('utf-8')).hexdigest()
        ozow_url = f"https://pay.ozow.com/?siteCode={site_code}&BankReference={bank_reference}&Amount={amount:.2f}&CurrencyCode=ZAR&Hash={generated_hash}"
        execute_ledger_split(amount)
        if get_session_user()["role"] == "customer":
            SYSTEM_DB["users"]["cust_1"]["tier"] = tier_upgrade
        session["cart"] = []
        return jsonify({"status": "redirect", "url": ozow_url})

    # Fallback to direct synchronous execution for mock or banking simulations
    execute_ledger_split(amount)
    if get_session_user()["role"] == "customer":
        SYSTEM_DB["users"]["cust_1"]["tier"] = tier_upgrade
    
    SYSTEM_DB["transactions"].append({
        "tx_id": tx_id, "amount": amount, "gateway": gateway, "timestamp": datetime.datetime.now().isoformat()
    })
    session["cart"] = []
    return jsonify({"status": "success", "message": f"Payment successfully processed via {gateway.upper()} interface pipeline.", "tx_id": tx_id})

# --- Ingestion File Stream Data Processing Pipeline ---
@app.route('/api/upload', methods=['POST'])
def api_upload():
    user = get_session_user()
    cv_text = request.form.get("cv_text", "").strip()
    job_desc = request.form.get("job_description", "").strip()
    
    if 'cv_file' in request.files:
        uploaded_file = request.files['cv_file']
        if uploaded_file.filename != '':
            file_extension = os.path.splitext(uploaded_file.filename)[1].lower()
            file_bytes = uploaded_file.read()
            
            if file_extension == '.pdf':
                cv_text = extract_text_from_pdf(file_bytes)
            elif file_extension in ['.txt', '.csv']:
                cv_text = file_bytes.decode('utf-8', errors='ignore').strip()
            elif file_extension == '.docx':
                try:
                    doc = Document(io.BytesIO(file_bytes))
                    cv_text = "\n".join([p.text for p in doc.paragraphs]).strip()
                except Exception as e:
                    return jsonify({"status": "error", "message": f"Word Document parser issue: {str(e)}"}), 400

    if not cv_text or not job_desc:
        return jsonify({"status": "error", "message": "Please supply full text profiles or file sheets for both input sections."}), 400
        
    user["cv_text"] = cv_text
    user["job_description"] = job_desc
    return jsonify({"status": "success", "message": "Application documents verified in local context session."})

# --- Complete Analysis Suite Engine Node ---
@app.route('/api/analyze', methods=['POST'])
def api_analyze():
    user = get_session_user()
    if not user["cv_text"] or not user["job_description"]:
        return jsonify({"status": "error", "message": "No active document matrix found. Upload variables first."}), 400
        
    # Free Tier Constraint Check Framework
    if user["tier"] == "free":
        if user["usage_count"] >= 1:
            return jsonify({"status": "error", "message": "Free tier capacity limit reached. Upgrade to Basic, Pro, or Enterprise to run unlimited cycles."}), 403
        user["usage_count"] += 1

    prompt = f"""
    You are an expert recruitment matrix operating within 2026 corporate hiring realities.
    Analyze the provided CV text relative to the job specification guidelines.
    
    Candidate CV:
    {user['cv_text']}
    
    Target Requirements:
    {user['job_description']}
    
    Return exclusively a verified JSON structure matching this exact design structure layout:
    {{
       "ats_score": 89,
       "optimized_cv": "Full professional tailored human-readable resume content context...",
       "cover_letter": "A compelling, warm, persuasive cover letter matching job attributes...",
       "qa_prep": [
          {{"question": "Core structural performance tracking question?", "answer": "Strategic professional solution answer response formula..."}}
       ]
    }}
    Do not add text around the JSON object. Keep it clean.
    """
    
    try:
        if openai.api_key == "mock-key":
            mock_res = {
                "ats_score": 94,
                "optimized_cv": f"=== OPTIMIZED RESUME RECORD ===\n\nProfile Summary:\nHighly efficient systems deployment expert with comprehensive domain mastery matching the execution specs of the target position.\n\nDemonstrated Project Competencies:\n- Executed target functions aligned directly to metrics in: {user['job_description'][:100]}...",
                "cover_letter": "Dear Hiring Matrix Core,\n\nI write with professional enthusiasm to present my qualifications for the open target specification role. My background in structural optimization makes me a solid match...",
                "qa_prep": [
                    {"question": "How do you align fast-moving development architectures with local constraints?", "answer": "By validating parameters against modular checks and early deployment pipeline testing routines."}
                ]
            }
            user.update(mock_res)
            return jsonify(mock_res)

        response = openai.chat.completions.create(
            model="gpt-4o", response_format={"type": "json_object"},
            messages=[{"role": "user", "content": prompt}], temperature=0.3
        )
        parsed = json.loads(response.choices[0].message.content)
        user.update(parsed)
        return jsonify(parsed)
    except Exception as e:
        return jsonify({"status": "error", "message": f"Analysis engine error: {str(e)}"}), 500

# --- Interactive 2026 AI Live Tutor Engine ---
@app.route('/api/tutor/chat', methods=['POST'])
def api_tutor_chat():
    user = get_session_user()
    
    if user["tier"] == "free":
        return jsonify({"response": "The Live AI Interview Coach requires an active system subscription (Pro, Enterprise, or AI Tutor standalone module Tier). Head to the platform store to activate."})
        
    msg = request.json.get("message", "").strip()
    if not msg:
        return jsonify({"status": "error", "message": "Empty query strings invalid."}), 400
        
    tutor_prompt = f"""
    You are an expert corporate executive interviewer running a live technical interview in 2026. 
    The candidate is applying for the job specified here: {user.get('job_description', 'General Developer')}
    Their professional background matrix context: {user.get('cv_text', 'Technical Specialist')}
    
    Do not mock or simulate nicely—ask realistic, sharp, professional behavioral questions. 
    Analyze their answer text for errors or weaknesses, provide direct constructive feedback, and then ask the next realistic progression question.
    """
    try:
        if openai.api_key == "mock-key":
            return jsonify({"response": "That answer focuses on standard execution metrics. To sound fully authoritative in a 2026 landscape, you need to state how you managed data safe-harbor compliance loop verification. Let's trace that out—how did you confirm localized tracking metrics remained secure during server optimization routines?"})
            
        res = openai.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": tutor_prompt}, {"role": "user", "content": msg}],
            temperature=0.6
        )
        return jsonify({"response": res.choices[0].message.content})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# --- Live Production Multi-Format Export Compiler Engine ---
@app.route('/api/download/<fmt>')
def download_compiled_file(fmt):
    user = get_session_user()
    cv_data = user.get("optimized_cv", "Run platform optimization cycle first.")
    cl_data = user.get("cover_letter", "Run platform optimization cycle first.")
    qa_data = "\n\n".join([f"Q: {i['question']}\nA: {i['answer']}" for i in user.get("qa_prep", [])])
    
    full_manifest_text = f"STORMCORE SYSTEM PLATFORM EXPORT FILE\n\n=== RESUME REWRITE ===\n{cv_data}\n\n=== COVER LETTER ===\n{cl_data}\n\n=== PRACTICE MATERIAL ===\n{qa_data}"
    
    if fmt == "txt":
        return send_file(io.BytesIO(full_manifest_text.encode('utf-8')), mimetype='text/plain', as_attachment=True, download_name='Stormcore_Application_Pack.txt')
    elif fmt == "csv":
        csv_content = f"Document Section,Content Data\n\"Optimized CV\",\"{cv_data.replace('&quot;', '\"')}\"\n\"Cover Letter\",\"{cl_data.replace('&quot;', '\"')}\""
        return send_file(io.BytesIO(csv_content.encode('utf-8')), mimetype='text/csv', as_attachment=True, download_name='Stormcore_Data_Manifest.csv')
    elif fmt == "pdf":
        return send_file(io.BytesIO(generate_pdf_bytes("Stormcore Application Deliverable Pack", full_manifest_text)), mimetype='application/pdf', as_attachment=True, download_name='Stormcore_Application_Pack.pdf')
    elif fmt == "docx":
        return send_file(io.BytesIO(generate_docx_bytes("Stormcore Application Deliverable Pack", full_manifest_text)), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='Stormcore_Application_Pack.docx')
    elif fmt == "zip":
        z_buf = io.BytesIO()
        with zipfile.ZipFile(z_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("01_Optimized_Resume.docx", generate_docx_bytes("Optimized Resume Rewrite", cv_data))
            zf.writestr("02_Custom_Cover_Letter.docx", generate_docx_bytes("Custom Cover Letter", cl_data))
            zf.writestr("03_Interview_Preparation_Guide.pdf", generate_pdf_bytes("Interview Preparation Blueprint", qa_data))
            zf.writestr("04_Raw_Data_Matrix.csv", f"Field,Value\nScore,{user.get('ats_score',0)}".encode('utf-8'))
        z_buf.seek(0)
        return send_file(z_buf, mimetype='application/zip', as_attachment=True, download_name='Stormcore_Master_Bundle.zip')
        
    return jsonify({"status": "error", "message": "Invalid download parameters specified."}), 400

# --- Intelligent Marketing Engine & Dynamic Platform Ad Algorithms ---
@app.route('/api/marketing/generate', methods=['POST'])
def api_generate_ads():
    platform = request.json.get("platform", "linkedin")  # facebook, linkedin, twitter, google
    tier = request.json.get("tier", "basic")
    
    # Live Programmatic 2026 Ad Policies Check Matrix
    compliance_rules = {
        "linkedin": "Prohibit programmatic automated spam outreach wording. Require verifiable identity credential terms. Enforce safe career progress representation.",
        "facebook": "Strict restriction on micro-targeting vulnerable employment classes without balanced disclosure. Ban deceptive income-generation click loops.",
        "google": "Requires clear landing architecture with transparent, accessible pricing frames. Zero tolerance for artificial dynamic ticking countdown clocks or single-day false scarcity claims."
    }
    
    rule_spec = compliance_rules.get(platform, "Maintain high-trust representation guidelines.")
    
    prompt = f"""
    You are an automated ad targeting specialist algorithm. 
    Generate high-converting, premium ad copy and landing page headers focused on the rules and frameworks below.
    
    Target Framework Level: {tier.upper()} Tier Setup.
    Platform Destination Matrix: {platform.upper()}
    Legal Network Compliance Restriction: {rule_spec}
    South African Consumer Protection Act Reference: Act 68 of 2008 Clear Disclosure Protocol.
    Global Regulation: Safe Harbors & Clear Value Discovery Pricing Framework.
    
    Create an ad layout tailored perfectly to attract impulse buyers searching for immediate solution tools.
    Provide output format structure inside clean JSON:
    {{
       "ad_headline": "Headline copy text matching compliance boundaries...",
       "ad_body_text": "High impact, value-driven text body targeted at converting action clicks immediately...",
       "landing_page_title": "Direct value statement title for the lander matching this layout tier..."
    }}
    Do not add extra formatting parameters. Valid JSON layout only.
    """
    try:
        if openai.api_key == "mock-key":
            return jsonify({
                "ad_headline": f"Secure Your Next Professional Step with Verified 2026 Tracking Standards [Platform Verified]",
                "ad_body_text": f"Stop allowing hidden corporate applicant parsers to filter out your real background. Unlock a clear, fully rewritten text profile matching modern industry specifications instantly.",
                "landing_page_title": f"Stormcore-SI Career Accelerator Core Framework - Welcome to the {tier.capitalize()} Matrix Tool Suite"
            })
            
        res = openai.chat.completions.create(
            model="gpt-4o", response_format={"type": "json_object"},
            messages=[{"role": "user", "content": prompt}], temperature=0.5
        )
        return jsonify(json.loads(res.choices[0].message.content))
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
